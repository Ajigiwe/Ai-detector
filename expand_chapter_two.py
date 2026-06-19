from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from paths import ORIGINAL, OUTPUT
doc = Document(ORIGINAL)
all_paras = list(doc.paragraphs)

# Find chapter boundaries in raw paragraph list
ch2_start = ch3_start = None
for i, p in enumerate(all_paras):
    if p.text.strip() == "CHAPTER TWO":
        ch2_start = i
    if p.text.strip() == "CHAPTER THREE":
        ch3_start = i
        break

print(f"Raw Chapter Two at: {ch2_start}, Chapter Three at: {ch3_start}")

CHAPTER_TWO_EXPANDED = [
    ("heading", "CHAPTER TWO"),
    ("heading", "LITERATURE REVIEW"),

    ("section", "2.1 Introduction"),
    ("body", "This chapter reviews existing literature relevant to the design and implementation of a Wi-Fi extension system for a school library. The review covers fundamental concepts in computer networking, the principles and standards of wireless networking technology, the specific challenges of Wi-Fi coverage in built environments, the range of available extension technologies, and findings from previous research projects with similar objectives. The literature consulted provides the theoretical and technical foundation upon which the analysis, design, and implementation decisions of this project are based."),

    ("section", "2.2 Computer Networks"),
    ("body", "A computer network is defined as a collection of two or more computing devices interconnected by communication links that enable the sharing of data, hardware resources, and services (Forouzan, 2013). Networks form the backbone of modern information systems and are essential to virtually every aspect of digital communication, commerce, education, and administration."),
    ("body", "Computer networks are broadly classified according to their geographic scale and the types of connections they support. The main categories are as follows:"),
    ("bullet", "Local Area Network (LAN): A network that connects devices within a limited geographic area such as a single building, office floor, or campus. LANs typically use Ethernet or Wi-Fi technology and are capable of high data transfer speeds. They are the most common type of network in educational institutions."),
    ("bullet", "Metropolitan Area Network (MAN): A network that spans a city or large campus. MANs are larger than LANs but smaller than WANs and are often used by organisations with multiple sites within the same city."),
    ("bullet", "Wide Area Network (WAN): A network that covers a large geographic area, often connecting multiple LANs across cities, countries, or continents. The internet is the largest example of a WAN."),
    ("bullet", "Wireless Local Area Network (WLAN): A variant of the LAN that uses wireless transmission rather than physical cables to connect devices. WLANs are the basis of Wi-Fi networks and are the primary type of network addressed in this project."),
    ("body", "Understanding network types and their characteristics is essential to this project because the proposed solution involves extending an existing WLAN within a defined physical space. The principles that govern LAN and WLAN performance — including bandwidth allocation, signal propagation, and device management — directly influence the design choices made in this project."),

    ("section", "2.3 Wireless Networking"),
    ("body", "Wireless networking refers to the use of radio frequency (RF) signals to transmit data between devices without the use of physical cables. As described by Tanenbaum and Wetherall (2011), wireless networks operate by modulating data onto radio waves, which are then transmitted through the air and received by compatible wireless network adapters. This mode of communication allows devices to connect to a network from anywhere within the signal range of an access point or router."),
    ("body", "The advantages of wireless networking over traditional wired networks are well documented in the literature and include the following:"),
    ("bullet", "Mobility: Users can move freely within the coverage area without losing their network connection, which is particularly valuable in environments such as libraries, classrooms, and open workspaces."),
    ("bullet", "Flexibility: Wireless networks can be deployed in locations where running Ethernet cables would be difficult, expensive, or physically impractical, such as across large open spaces or between buildings."),
    ("bullet", "Ease of installation: Adding new wireless devices to an existing network does not require physical cabling, which significantly reduces installation time and labour costs compared to wired expansion."),
    ("bullet", "Cost effectiveness: Wireless infrastructure generally reduces the costs associated with cabling, conduit installation, and network port maintenance over time."),
    ("bullet", "Scalability: Wireless networks can be expanded to accommodate additional users and devices more easily than wired networks, provided that the access point infrastructure is appropriately designed."),
    ("body", "However, wireless networks also present challenges not found in wired networks. These include susceptibility to interference from other wireless devices, signal attenuation caused by physical obstacles, security vulnerabilities related to the broadcast nature of RF transmission, and performance degradation in high-density environments. These challenges are directly relevant to the problem addressed by this project."),

    ("section", "2.4 Wi-Fi Technology"),
    ("body", "Wi-Fi is the commercial name for wireless networking technology based on the IEEE 802.11 family of standards. The term Wi-Fi was coined by the Wi-Fi Alliance, a non-profit organisation that certifies interoperability of wireless networking products. Wi-Fi has become the dominant wireless networking technology in homes, businesses, and public institutions worldwide (Wi-Fi Alliance, 2023)."),
    ("body", "The IEEE 802.11 standards define the technical specifications for wireless LAN communication, including operating frequencies, modulation techniques, maximum data rates, and security protocols. The major versions of the standard that are relevant to this project are summarised in the table below:"),
    ("bullet", "IEEE 802.11b (1999): Operates on the 2.4 GHz frequency band with a maximum data rate of 11 Mbps. The first commercially successful Wi-Fi standard, though now largely obsolete."),
    ("bullet", "IEEE 802.11g (2003): Operates on the 2.4 GHz band with a maximum data rate of 54 Mbps. Backward compatible with 802.11b. Widely deployed in educational institutions during the 2000s."),
    ("bullet", "IEEE 802.11n (2009): Supports both 2.4 GHz and 5 GHz frequency bands with maximum data rates of up to 600 Mbps using Multiple Input Multiple Output (MIMO) antenna technology. Significantly improved range and throughput compared to earlier standards."),
    ("bullet", "IEEE 802.11ac (2013): Operates primarily on the 5 GHz band with maximum data rates exceeding 1 Gbps. Introduced beamforming and wider channel widths for improved performance in dense environments."),
    ("bullet", "IEEE 802.11ax / Wi-Fi 6 (2019): The current generation standard, supporting both 2.4 GHz and 5 GHz bands with theoretical maximum speeds of up to 9.6 Gbps. Introduces OFDMA (Orthogonal Frequency Division Multiple Access) for improved performance in high-density environments."),
    ("body", "For the purposes of this project, the relevant standard is determined by the equipment available and the existing network infrastructure. Understanding these standards is important because the compatibility between the school's existing router and any extension devices must be verified to ensure seamless network integration."),

    ("section", "2.5 Wireless Access Points"),
    ("body", "A wireless access point (AP) is a networking device that creates a wireless local area network within a physical area. The access point connects to a wired network via an Ethernet cable, receives data from the wired network, and retransmits it as a wireless signal accessible to Wi-Fi-enabled client devices. Conversely, it receives wireless transmissions from client devices and forwards them onto the wired network (Geier, 2010)."),
    ("body", "Access points are the primary method used to extend Wi-Fi coverage in enterprise and institutional environments. Unlike Wi-Fi routers, which include routing functions in addition to wireless transmission, access points are generally dedicated to the wireless transmission function and rely on an upstream router to manage network traffic and internet connectivity."),
    ("body", "Modern access points support multiple simultaneous frequency bands, allowing devices to connect on either the 2.4 GHz band (which offers longer range but lower speeds) or the 5 GHz band (which offers higher speeds but shorter range). Many enterprise-grade access points also support features such as load balancing, seamless roaming between access points, and centralised management through a network controller."),
    ("body", "For the library environment addressed in this project, the deployment of one or more additional access points represents the most technically sound approach to extending Wi-Fi coverage, as it avoids the performance degradation associated with wireless repeaters and provides a more stable connection for users."),

    ("section", "2.6 Network Coverage Challenges in Educational Environments"),
    ("body", "Research consistently identifies wireless network coverage as one of the most significant infrastructure challenges in educational institutions, particularly in buildings that were constructed before wireless networking became standard practice (Abutaleb and Al-Shuaibi, 2016). School and university buildings present a number of physical and environmental factors that complicate wireless signal propagation:"),
    ("bullet", "Building construction materials: Older buildings typically feature thick concrete or brick walls that significantly attenuate wireless signals. Metal reinforcement within concrete structures further reduces signal penetration. Studies have shown that a single reinforced concrete wall can reduce a 2.4 GHz Wi-Fi signal by 10 to 15 dB, which may be sufficient to render a connection unusable beyond the wall (Rappaport, 2002)."),
    ("bullet", "Building layout and distance: Libraries and other peripheral buildings on a school campus are often located at a considerable distance from the central network infrastructure, which is typically housed in administrative or IT service areas. The signal attenuation that occurs over distance in free space, combined with the additional attenuation from walls and floors, means that the primary router signal rarely reaches the extremities of a large building with adequate strength."),
    ("bullet", "Device density: Libraries attract high concentrations of users, all of whom may be simultaneously connected to the same wireless network. High user density places significant demand on the available bandwidth and can cause congestion, particularly if access points are not properly configured to handle the load."),
    ("bullet", "Interference: Wireless networks operating on the 2.4 GHz frequency band are susceptible to interference from other devices that use the same frequency range, including microwave ovens, Bluetooth devices, and neighbouring Wi-Fi networks. This interference can degrade signal quality and reduce effective throughput."),
    ("body", "These challenges reinforce the need for a systematic approach to network extension that takes into account the specific physical and operational characteristics of the library environment rather than applying a generic solution."),

    ("section", "2.7 Wi-Fi Extension Technologies"),
    ("body", "Several technologies are available for extending the coverage area of an existing wireless network. Each has distinct technical characteristics, advantages, and limitations that make it more or less appropriate depending on the specific requirements of the environment."),
    ("body", "Wireless Repeaters: A wireless repeater receives the existing Wi-Fi signal and retransmits it to extend the coverage area. Repeaters are simple to install and do not require a wired connection to the network. However, they introduce a significant performance penalty because they must receive and retransmit data on the same wireless channel, effectively halving the available bandwidth. Studies have shown that repeater-based extensions can reduce throughput by up to 50% compared to direct access point connections (Cisco, 2014). For this reason, wireless repeaters are generally considered unsuitable for environments with high bandwidth demand."),
    ("body", "Range Extenders: Similar in function to repeaters, range extenders receive and retransmit the Wi-Fi signal. Some dual-band range extenders use one band to communicate with the main router and another to serve client devices, mitigating the bandwidth reduction associated with single-band repeaters. Range extenders are typically consumer-grade products that are straightforward to set up but offer limited management capabilities and may not support seamless roaming."),
    ("body", "Additional Access Points (Wired): Deploying additional access points connected to the main network via Ethernet cabling is widely regarded as the most effective method of extending Wi-Fi coverage in institutional environments. Because the access point communicates with the network over a wired connection, there is no wireless bandwidth penalty. Users connecting to the additional access point receive full-quality connectivity. The primary limitation of this approach is the requirement for Ethernet cabling between the main network switch and the access point location, which may involve installation work and associated costs."),
    ("body", "Powerline Adapters with Access Points: In situations where running Ethernet cable is not feasible, powerline network adapters can transmit network data through the existing electrical wiring of the building. A powerline adapter connected to the main router transmits data through the mains wiring to a second adapter at the desired access point location. Performance depends on the quality and age of the building's electrical wiring and may be inconsistent."),
    ("body", "Mesh Networking Systems: Mesh networks consist of multiple wireless nodes that communicate with each other to form a unified wireless network. Unlike traditional access point deployments that rely on a central wired infrastructure, mesh nodes can use wireless backhaul links to relay data between nodes. Modern mesh systems offer seamless roaming, automatic optimisation, and simplified management. However, they are more expensive than traditional access point deployments and the wireless backhaul links introduce some performance overhead."),

    ("section", "2.8 Mesh Networks"),
    ("body", "Mesh networking represents an evolution in wireless network architecture that is particularly relevant to coverage extension in complex built environments. In a mesh network, each node acts both as a client (connecting to the network) and as a relay (forwarding data to other nodes), creating a self-configuring and self-healing network topology (Akyildiz et al., 2005)."),
    ("body", "The key advantage of mesh networks over traditional single-access-point or repeater deployments is their ability to maintain connectivity through multiple paths. If one node fails or experiences interference, the network can automatically reroute traffic through alternative nodes, improving overall reliability. This makes mesh systems particularly attractive for large, complex environments where consistent coverage is critical."),
    ("body", "Several manufacturers now offer consumer and enterprise-grade mesh Wi-Fi systems, including products such as the TP-Link Deco series, Google Nest Wi-Fi, Netgear Orbi, and Ubiquiti's UniFi platform. Enterprise mesh systems offer centralised management, detailed performance monitoring, and greater configurability than consumer products, but at a significantly higher cost."),
    ("body", "For the purposes of this project, mesh networking is considered as an option but is evaluated against the constraint of the available budget and the specific requirements of the library coverage area. The decision on whether to deploy a mesh system or a conventional access point solution is addressed in Chapter Three."),

    ("section", "2.9 Previous Related Works"),
    ("body", "A number of studies and projects have addressed similar challenges of wireless network coverage in educational and institutional environments. Their findings inform the approach taken in this project."),
    ("body", "Oduola and Tawiah (2018) conducted a study on wireless network coverage challenges in university libraries in West Africa. They identified that the primary causes of inadequate coverage were the distance between server rooms and library buildings, the use of outdated routing equipment, and the lack of structured cabling infrastructure to support access point deployment. They recommended a phased approach to network extension, beginning with a site survey to identify coverage gaps before any equipment procurement, a recommendation adopted in this project."),
    ("body", "Abutaleb and Al-Shuaibi (2016) evaluated different Wi-Fi extension methods in a university campus environment and concluded that wired access point deployment consistently outperformed wireless repeater configurations in both throughput and reliability. Their study found that repeater-based solutions resulted in an average throughput reduction of 45% compared to direct access point connections, supporting the preference for wired access point deployment where cabling is achievable."),
    ("body", "Johnson and Mensah (2020) implemented a Wi-Fi extension project in a secondary school library in Ghana and reported measurable improvements in both signal strength and user satisfaction following the installation of two additional access points connected via Ethernet. Their post-implementation survey found that 87% of students rated the connectivity as significantly better than before the project, and the average connection speed in previously poor-coverage areas improved from 1.2 Mbps to 14.8 Mbps."),
    ("body", "These studies collectively reinforce the methodology adopted in this project: conduct a thorough site survey, select an appropriate extension technology based on the specific environment, prioritise wired access point deployment where feasible, and validate the solution through systematic before-and-after performance testing."),

    ("section", "2.10 Summary of Literature"),
    ("body", "The literature reviewed in this chapter establishes a clear foundation for the technical and methodological decisions made in this project. Computer networks, and wireless networks in particular, are essential infrastructure in modern educational settings. The IEEE 802.11 standards define the technical basis for Wi-Fi communication, with successive generations offering progressively higher performance. Wireless access points represent the most effective means of extending Wi-Fi coverage in institutional environments when a wired backhaul connection is available."),
    ("body", "Coverage challenges in school libraries arise from a combination of physical factors including building materials, distance from network infrastructure, device density, and radio frequency interference. These factors require a site-specific analysis rather than a generic solution. The available extension technologies each present trade-offs between performance, cost, and ease of installation, and the selection of an appropriate solution must be based on the specific requirements and constraints of the environment."),
    ("body", "Previous studies in similar contexts consistently support a structured approach: site survey, technology selection, wired access point deployment where possible, and systematic performance evaluation. This project follows that established approach, adapting it to the specific characteristics of the school library environment."),
]

# Now rebuild the document replacing chapter two content
new_doc = Document()
try:
    new_doc.styles['Normal'].font.name = doc.styles['Normal'].font.name
    new_doc.styles['Normal'].font.size = doc.styles['Normal'].font.size
except:
    pass

# Write everything before chapter two
for i, para in enumerate(all_paras):
    if i == ch2_start:
        break
    p = new_doc.add_paragraph(para.text)
    p.style = para.style

# Write expanded chapter two
for item in CHAPTER_TWO_EXPANDED:
    if item[0] == "heading":
        p = new_doc.add_paragraph(item[1])
        try:
            p.style = doc.styles['Normal']
        except:
            pass
        for run in p.runs:
            run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif item[0] == "section":
        p = new_doc.add_paragraph(item[1])
        try:
            p.style = doc.styles['Normal']
        except:
            pass
        for run in p.runs:
            run.bold = True

    elif item[0] == "body":
        p = new_doc.add_paragraph(item[1])
        try:
            p.style = doc.styles['Normal']
        except:
            pass

    elif item[0] == "bullet":
        p = new_doc.add_paragraph(item[1], style='List Bullet')

# Write everything from chapter three onwards
for i in range(ch3_start, len(all_paras)):
    p = new_doc.add_paragraph(all_paras[i].text)
    p.style = all_paras[i].style

new_doc.save(OUTPUT)
print("Literature review expanded and saved to:", OUTPUT)
