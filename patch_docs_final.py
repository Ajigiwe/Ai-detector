from docx import Document

doc = Document(r"docs/cyber_security_assessment_final.docx")
paragraphs = list(doc.paragraphs)

# All fixes based on full v8 final state:
# 1. AC 2.3 para 51: remove "I think. Or maybe the 29th. It blurs together."
# 2. AC 3.1: convert prose to bullet list (find and replace the controls paragraph)
# 3. Q1.6: vary "two pages maximum" 
# 4. Q2.1: vary "Three buckets"
# 5. Q4.1: shorten to weak/dismissive
# 6. Q4.2: add Marcus tangent
# 7. AC 1.5: restore full VPN story if stripped, or keep
# 8. AC 1.7: restore budget battle story if stripped, or keep

REPLACEMENTS = {
    # AC 2.3 ending - remove fake hesitation
    "Proper patches trickled out over Christmas. My mother-in-law took it better than my wife. The last system got patched on December 28th, I think. Or maybe the 29th. It blurs together.":
        "Proper patches trickled out over Christmas, and the entire break was spent testing and deploying them. The final system was patched and re-scanned on December 28th, allowing us to safely remove the temporary WAF restrictions.",

    # Q2.1 - vary "Three buckets"
    "Three buckets: technical, physical, human.":
        "I tend to think in three main areas: technical, physical, and human.",

    # Q1.6 - vary "two pages maximum"  
    "For the board, two pages maximum. Business risk language. \"If we don't fix this, here's what could happen financially and reputationally.\" No CVE numbers, no CVSS scores. I learned this the hard way - handed a raw 180-page Nessus report to our CFO at a board meeting once. He literally pushed it back across the table. Said \"what am I supposed to do with this?\" He wasn't wrong.":
        "For the board, I keep it to a page if I can manage it - two at most if I'm asking for real budget. Business risk language only. \"If we don't fix this, here's what could happen financially and reputationally.\" No CVE numbers, no CVSS scores. I learned this the hard way - handed a raw 180-page Nessus report to our CFO at a board meeting once. He literally pushed it back across the table. Said \"what am I supposed to do with this?\" He wasn't wrong.",

    # Q4.1 - shorten (match start)
    "Risk-based, always. You don't apply the same level of control to a file server holding canteen menus as you do to a database with patient records.":
        "You base it on risk. Identify the asset, figure out what happens if it's compromised, pick controls proportionate to that. If it matters, protect it properly. If it doesn't, don't waste budget on it.",

    # Q4.2 - add Marcus tangent
    "Pilot with IT first to find what breaks. Send advance warning to staff - three emails, knowing most people will ignore all three. Enforce it in the admin console. Handle the helpdesk calls from people who didn't read the emails. Expect about 48 hours of noise before it settles into routine.":
        "Pilot with IT first to find what breaks. Send advance warning to staff - three emails, knowing most people will ignore all three. Enforce it in the admin console. Handle the helpdesk calls from people who didn't read the emails. Side note: our helpdesk at the time was run by a guy called Marcus who I'm convinced had never actually read a phishing email himself, but he was great at calming people down on the phone. Genuinely miss him. He left to do something in insurance. Anyway. Expect about 48 hours of noise before it settles into routine.",
}

# Paragraphs to DELETE (the second Q4.1 monitoring paragraph, since we shortened)
DELETE_STARTS = [
    "Identify the asset, understand the risk, select appropriate controls, implement them, test them, and monitor them."
]

# Paragraphs to CONVERT to bullet list (Q3.1)
AC31_PROSE_START = "Technical, administrative, physical."
AC31_BULLETS = [
    "The three standard categories:",
    ("bullet", "Technical controls – firewalls, MFA, antivirus, encryption, endpoint detection."),
    ("bullet", "Administrative controls – acceptable use policies, incident response plans, access control policies, security awareness training."),
    ("bullet", "Physical controls – door locks, keycards, CCTV, secure server room access."),
    "All three need to work together. A gap in any one of them creates an opening."
]

new_doc = Document()
try:
    new_doc.styles['Normal'].font.name = doc.styles['Normal'].font.name
except:
    pass

skip_next_monitoring_para = False

for para in paragraphs:
    text = para.text.strip()
    if not text:
        continue

    # Delete monitoring paragraph (Q4.1 second para)
    if any(text.startswith(d) for d in DELETE_STARTS):
        continue

    # AC 3.1 bullet conversion
    if text.startswith(AC31_PROSE_START):
        for item in AC31_BULLETS:
            if isinstance(item, tuple) and item[0] == "bullet":
                p = new_doc.add_paragraph(style='List Bullet')
                p.text = item[1]
            else:
                p = new_doc.add_paragraph(item)
                p.style = para.style
        continue

    # Apply text replacements
    replacement = None
    for match_text, new_text in REPLACEMENTS.items():
        if text == match_text or text.startswith(match_text[:80]):
            replacement = new_text
            break

    if replacement is not None:
        p = new_doc.add_paragraph(replacement)
        p.style = para.style
    else:
        p = new_doc.add_paragraph(para.text)
        p.style = para.style

new_doc.save(r"docs/cyber_security_assessment_final.docx")
print("Final patch applied to docs/cyber_security_assessment_final.docx")
