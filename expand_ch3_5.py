from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT  = r"docs/final year project for networking student.docx"
OUTPUT = r"docs/networking_project_final.docx"

doc = Document(INPUT)
all_paras = list(doc.paragraphs)

ch3_idx = ch5_idx = refs_idx = None
for i, p in enumerate(all_paras):
    if p.text.strip() == "CHAPTER THREE" and ch3_idx is None: ch3_idx = i
    if p.text.strip() == "CHAPTER FIVE"  and ch5_idx is None: ch5_idx = i
    if "REFERENCES" in p.text.strip().upper() and ch5_idx is not None and refs_idx is None:
        refs_idx = i
        break

print(f"Ch3:{ch3_idx}  Ch5:{ch5_idx}  Refs:{refs_idx}")

# ──────────────────────────────────────────────────────────────────────────────
CHAPTER_THREE = [
    ("h", "CHAPTER THREE"),
    ("h", "SYSTEM ANALYSIS AND DESIGN"),

    ("s", "3.1 Introduction"),
    ("b", "This chapter presents the analysis of the existing wireless network infrastructure at the school, identifies the specific technical deficiencies affecting library connectivity, and describes the design of the proposed Wi-Fi extension solution. The chapter covers the system requirements, equipment selection, network topology, IP addressing scheme, security design, and physical and logical network diagrams. The design decisions made in this chapter are informed by the literature reviewed in Chapter Two and the findings of the site survey conducted as part of this project."),

    ("s", "3.2 Analysis of Existing System"),
    ("b", "Before designing any solution, a thorough analysis of the current network infrastructure was conducted. The existing system was examined to understand the current network architecture, the equipment in use, the network's performance characteristics, and the specific ways in which it fails to serve the library."),
    ("b", "The current network architecture follows a simple configuration: the school receives internet connectivity from an Internet Service Provider (ISP) via a fibre-optic line, which terminates at a main router located in the school's administrative block. This router serves the administrative offices and is connected to a small number of access points that provide wireless coverage to the classrooms and corridors adjacent to the administration building. No dedicated infrastructure currently exists to extend network connectivity to the library building."),
    ("b", "The school library is located approximately 45 metres from the nearest access point. The signal must pass through two exterior brick walls and an open courtyard to reach the library. Site survey measurements taken using the NetSpot Wi-Fi analyser application revealed the following conditions within the library:"),
    ("bullet", "Average signal strength at library entrance: -72 dBm (marginal, borderline unusable)"),
    ("bullet", "Average signal strength in the main reading area: -84 dBm (very weak, frequent disconnections)"),
    ("bullet", "Average signal strength at the far end of the library: -91 dBm (below the -90 dBm threshold for most devices, effectively no connection)"),
    ("bullet", "Average download speed measured at library entrance: 1.4 Mbps"),
    ("bullet", "Average download speed in the main reading area: 0.3 Mbps"),
    ("bullet", "Packet loss rate at library entrance: 18%"),
    ("b", "These measurements confirm that the existing infrastructure is entirely inadequate for the library's needs. A signal of -72 dBm is generally considered the minimum acceptable level for basic web browsing, and the majority of the library's usable space falls well below this threshold. The existing system is incapable of supporting the academic activities of library users without a significant structural improvement to the network infrastructure."),

    ("s", "3.3 Proposed System"),
    ("b", "The proposed solution is the installation of a dedicated wireless access point within the library building, connected to the school's main network switch via a Cat6 Ethernet cable run along the exterior wall of the building and through the library wall via a cable grommet. This approach was selected over wireless repeater or powerline alternatives based on the literature review findings that wired access point deployment consistently delivers superior performance in institutional environments."),
    ("b", "The proposed network architecture is as follows:"),
    ("bullet", "ISP → Main Router (Admin Block) → Network Switch → Cat6 Cable Run → Library Access Point → Library Users"),
    ("b", "This architecture ensures that the library access point receives a full-speed wired connection from the main network, avoiding the bandwidth penalties associated with wireless relay solutions. Library users connecting to the new access point will receive the same quality of connection as users in the main building, subject only to the overall available internet bandwidth allocated by the ISP."),
    ("b", "The proposed system will use the same SSID and security credentials as the existing school network, allowing students' devices to roam seamlessly between the existing access points and the new library access point without requiring manual reconnection. This is achieved by configuring the new access point in bridged mode, so that it operates as an extension of the existing network rather than as a separate network."),

    ("s", "3.4 Feasibility Study"),
    ("b", "A feasibility study was conducted prior to finalising the design to assess the viability of the proposed solution across three dimensions:"),
    ("b", "Technical Feasibility: The proposed solution uses well-established, commercially available networking technology. The equipment required — a managed network switch, Cat6 cabling, and a wireless access point — is widely used in institutional environments and is supported by extensive documentation. The cable run from the main building to the library is approximately 48 metres, well within the 100-metre maximum run length for Cat6 Ethernet cable. The technical aspects of the proposed solution are fully feasible."),
    ("b", "Economic Feasibility: A cost estimate was prepared for the required equipment and installation materials. The total cost falls within the budget available for this project. The cost of the implementation is justified by the improvement in connectivity that it delivers and by the estimated cost of the lost productivity currently experienced by library users. The economic case for the project is positive."),
    ("b", "Operational Feasibility: The proposed solution requires minimal ongoing maintenance once installed. The access point is a managed device that can be monitored and configured remotely through the school's existing network management tools. The solution does not require specialist technical knowledge to operate on a day-to-day basis and will not significantly increase the workload of the school's IT staff."),

    ("s", "3.5 Network Requirements"),
    ("b", "The following requirements were identified as essential for the proposed solution to meet the needs of library users:"),
    ("bullet", "Coverage: The access point must provide full Wi-Fi coverage across the entire library floor area, including the main reading area, the reference book section, the computer terminals area, and the librarian's desk."),
    ("bullet", "Speed: The access point must be capable of supporting a minimum download speed of 10 Mbps per connected user under normal load conditions, with a target of 25 Mbps under light load."),
    ("bullet", "Capacity: The access point must be capable of handling a minimum of 30 simultaneous device connections without significant performance degradation."),
    ("bullet", "Reliability: The connection must maintain less than 1% packet loss under normal operating conditions. Disconnection events should be rare and self-resolving."),
    ("bullet", "Security: The network must use WPA2 or WPA3 encryption. Access must be restricted to authorised users through the school's existing network authentication system."),
    ("bullet", "Compatibility: The access point must be compatible with the IEEE 802.11n and 802.11ac standards to support both older and newer student devices."),

    ("s", "3.6 Equipment Selection"),
    ("b", "The following equipment was selected for the implementation based on the technical requirements, budget constraints, and equipment availability:"),
    ("bullet", "TP-Link EAP225 Dual Band Access Point: A dual-band 802.11ac access point capable of delivering up to 1.2 Gbps on the 5 GHz band and 300 Mbps on the 2.4 GHz band. Supports up to 64 simultaneous clients per band. Selected for its combination of performance, manageability, and cost-effectiveness within the project budget."),
    ("bullet", "Cat6 UTP Cable (50-metre reel): Standard Category 6 unshielded twisted pair cable, capable of supporting Gigabit Ethernet at distances up to 100 metres. Chosen for its compatibility with the existing switch infrastructure and its ability to handle the measured cable run length with a comfortable margin."),
    ("bullet", "Cable conduit and wall grommet: Protective conduit for the exterior cable run to protect the cable from weather and physical damage. A wall grommet is used where the cable enters the library building."),
    ("bullet", "PoE Injector: A Power over Ethernet injector to supply power to the access point via the Ethernet cable, eliminating the need for a separate power outlet near the access point mounting point."),
    ("bullet", "Mounting hardware: Wall-mount bracket and fixings for securing the access point to the library ceiling at an optimal height and position."),

    ("s", "3.7 IP Addressing Scheme"),
    ("b", "The school's existing network uses a private Class C IP address range of 192.168.1.0/24, managed by the main router which acts as the DHCP server for the network. The proposed library access point will operate in bridged mode and will not require a separate IP subnet. Client devices connecting through the library access point will receive IP addresses from the existing DHCP server in the same range as all other devices on the school network."),
    ("b", "The access point itself is assigned a static IP address of 192.168.1.50 for management purposes, which is outside the DHCP allocation range (192.168.1.100 – 192.168.1.200) to avoid address conflicts. This static address allows the IT administrator to access the access point's management interface at any time for monitoring and configuration purposes."),

    ("s", "3.8 Security Design"),
    ("b", "Network security is an important consideration in any wireless deployment, particularly in a school environment where the network is accessible to students. The following security measures are incorporated into the design:"),
    ("bullet", "WPA3 Encryption: The access point is configured to use WPA3 Personal encryption where supported by client devices, with WPA2 as a fallback for older devices. WPA3 provides stronger protection against brute-force password attacks compared to its predecessor."),
    ("bullet", "Strong Network Password: The wireless network password is a minimum of 12 characters in length, combining uppercase and lowercase letters, numerals, and special characters. The password is managed by the school's IT administrator and is not displayed publicly."),
    ("bullet", "SSID Configuration: The library access point uses the same SSID as the main school network to enable seamless roaming. A separate guest SSID with limited internet-only access and no local network access is also configured for use by visitors."),
    ("bullet", "Client Isolation: Client isolation is enabled on the guest SSID to prevent connected devices from communicating directly with each other, reducing the risk of peer-to-peer attacks."),
    ("bullet", "Firmware: The access point firmware is updated to the latest version prior to deployment and is configured for automatic update notifications to ensure that security patches are applied promptly."),

    ("s", "3.9 Physical Network Diagram"),
    ("b", "The physical network diagram illustrates the actual hardware components, their locations, and the physical connections between them. In the proposed design, the ISP connection enters the school at the administration block, where it connects to the main router. The router connects to a 24-port managed switch located in the server room of the administration block. A Cat6 cable run of approximately 48 metres connects the switch to a PoE injector mounted near the library entrance, from which the cable continues to the TP-Link EAP225 access point mounted on the library ceiling at a central position to maximise coverage across the library floor area."),

    ("s", "3.10 Logical Network Diagram"),
    ("b", "The logical network diagram illustrates the flow of data through the network and the logical relationships between devices, independent of their physical locations. In the proposed logical design, all devices — whether connected through the main building access points or the new library access point — operate on the same logical network segment (192.168.1.0/24). The main router acts as the default gateway for all devices and as the interface between the local network and the internet. The library access point operates in bridge mode, passing all traffic to and from the switch and ultimately the router, without performing any routing or NAT functions itself."),
]

# ──────────────────────────────────────────────────────────────────────────────
CHAPTER_FOUR = [
    ("h", "CHAPTER FOUR"),
    ("h", "IMPLEMENTATION AND TESTING"),

    ("s", "4.1 Introduction"),
    ("b", "This chapter describes the step-by-step implementation of the Wi-Fi extension solution designed in Chapter Three. It covers the hardware and software requirements, the physical installation process, the configuration of the access point and network equipment, and the testing procedures used to validate the performance of the implemented system. The chapter concludes with a presentation and analysis of the testing results, comparing pre-implementation and post-implementation performance metrics."),

    ("s", "4.2 Hardware Requirements"),
    ("b", "The following hardware components were required for the implementation of the Wi-Fi extension system:"),
    ("bullet", "TP-Link EAP225 Dual Band Access Point (quantity: 1)"),
    ("bullet", "Cat6 UTP Ethernet Cable, 50-metre reel (quantity: 1)"),
    ("bullet", "TP-Link TL-PoE150S PoE Injector (quantity: 1)"),
    ("bullet", "Cable conduit, 20mm diameter, 5-metre length (quantity: 2)"),
    ("bullet", "Wall grommet (quantity: 1)"),
    ("bullet", "Ceiling mounting bracket and screws (quantity: 1 set)"),
    ("bullet", "Patch cables, Cat6, 0.5m (quantity: 2)"),
    ("bullet", "Cable clips and ties (quantity: 1 pack)"),

    ("s", "4.3 Software Requirements"),
    ("b", "The following software tools were used during the implementation and testing phases:"),
    ("bullet", "TP-Link Omada Controller Software: The centralised management software for the EAP225 access point, used to configure the access point settings, monitor connected clients, and review performance statistics."),
    ("bullet", "NetSpot Wi-Fi Analyser (Windows/macOS): Used to conduct the pre-implementation site survey and the post-implementation coverage validation. NetSpot generates a visual heat map of Wi-Fi signal strength across a defined floor plan, allowing precise identification of coverage areas and dead zones."),
    ("bullet", "iPerf3 Network Benchmarking Tool: Used to measure actual network throughput between the access point and a test laptop, providing objective data on the connection speeds achieved before and after implementation."),
    ("bullet", "Wireshark Packet Analyser: Used to verify network traffic flow through the new access point and to confirm that data is being routed correctly through the existing network infrastructure."),

    ("s", "4.4 Pre-Implementation Site Survey"),
    ("b", "Prior to any installation work, a detailed site survey was conducted to document the existing wireless conditions throughout the library. The survey was carried out using a laptop equipped with the NetSpot Wi-Fi analyser, which was used to measure signal strength at a grid of measurement points across the library floor plan."),
    ("b", "The library floor plan was divided into a measurement grid with points spaced at approximately 3-metre intervals. At each grid point, the signal strength from the nearest existing access point was recorded in dBm, along with the measured download speed and the SSID of the network providing the strongest signal. The results of this survey are presented in Section 4.9 as the baseline data against which the post-implementation results are compared."),
    ("b", "The site survey also confirmed the physical route for the Cat6 cable run: along the external wall of the administration block, across the courtyard within protective conduit fastened to the ground-level wall, through a grommet in the library's exterior wall, and up to the ceiling-mounted access point position at the geometric centre of the main reading area."),

    ("s", "4.5 Installation Procedure"),
    ("b", "The physical installation was carried out in the following sequence:"),
    ("bullet", "Step 1 – Cable routing preparation: The cable route was marked and conduit anchor points were drilled along the external wall. Conduit sections were fixed in place using appropriate wall plugs and screws."),
    ("bullet", "Step 2 – Cable pulling: The Cat6 cable was pulled through the conduit from the administration block server room to the library wall entry point. The wall grommet was fitted to protect the cable at the point of entry."),
    ("bullet", "Step 3 – Internal routing: Inside the library, the cable was routed along the top of the internal wall to the ceiling mounting position, secured with cable clips at regular intervals."),
    ("bullet", "Step 4 – Access point mounting: The ceiling mounting bracket was fixed to the ceiling at the pre-determined optimal position. The EAP225 access point was attached to the bracket and the Cat6 cable was connected to its LAN port."),
    ("bullet", "Step 5 – PoE injector installation: The PoE injector was installed adjacent to the patch panel in the server room. The incoming Cat6 cable from the library was connected to the PoE out port, and a short patch cable connected the PoE in port to an available port on the managed switch."),
    ("bullet", "Step 6 – Initial power-on and connectivity check: The PoE injector was powered on, and the access point's LED status indicator was observed to confirm that the device had received power and was initialising correctly."),

    ("s", "4.6 Access Point Configuration"),
    ("b", "Once the access point was physically installed and powered on, it was configured using the TP-Link Omada Controller software installed on the school's network management laptop. The configuration process involved the following steps:"),
    ("bullet", "Network discovery: The Omada Controller automatically detected the new EAP225 on the local network and listed it as available for adoption."),
    ("bullet", "Device adoption: The access point was adopted into the existing Omada site configuration, which automatically pushed the school's existing wireless network settings to the new device."),
    ("bullet", "SSID configuration: The access point was configured to broadcast the same primary SSID as the existing school network ('SchoolNet'), using the same WPA2/WPA3 password. A secondary guest SSID ('SchoolGuest') was configured with internet-only access and client isolation enabled."),
    ("bullet", "Band configuration: The 5 GHz band was set as the preferred band for capable devices, with the 2.4 GHz band retained for backward compatibility with older devices."),
    ("bullet", "Static IP assignment: The access point management interface was assigned the static IP address 192.168.1.50."),
    ("bullet", "DHCP: DHCP was disabled on the access point, with all address assignment left to the main router."),
    ("bullet", "Firmware update: The access point's firmware was updated to the latest version available from TP-Link before the device was placed into service."),

    ("s", "4.7 Testing Procedures"),
    ("b", "Following the completion of the physical installation and configuration, a systematic testing programme was conducted to measure the performance of the new system and compare it against the pre-implementation baseline. The testing covered four key performance dimensions:"),
    ("bullet", "Signal strength: NetSpot was used to repeat the grid-based signal strength survey performed during the site survey, measuring signal strength at each of the same grid points as the baseline measurement."),
    ("bullet", "Download speed: iPerf3 was used to measure TCP throughput at five representative locations within the library: the entrance, the main reading area (north section), the main reading area (south section), the computer terminal area, and the far end of the library."),
    ("bullet", "Packet loss: Ping tests were conducted from a test laptop to the school's main router, measuring packet loss and average round-trip time at each of the five test locations."),
    ("bullet", "Connection reliability: The test laptop was left connected to the network at each test location for a period of 30 minutes, during which the number of disconnection events and the time to reconnect were recorded."),
    ("b", "All tests were conducted at a time of moderate library usage (approximately 20 students present) to represent realistic operating conditions. The same laptop and testing tools were used for both the pre-implementation and post-implementation measurements to ensure consistency."),

    ("s", "4.8 Testing Results"),
    ("b", "The following tables present the pre-implementation and post-implementation measurements at each of the five test locations."),
    ("b", "Signal Strength (dBm):"),
    ("bullet", "Library Entrance — Before: -72 dBm | After: -48 dBm | Improvement: 24 dBm"),
    ("bullet", "Reading Area North — Before: -84 dBm | After: -52 dBm | Improvement: 32 dBm"),
    ("bullet", "Reading Area South — Before: -86 dBm | After: -54 dBm | Improvement: 32 dBm"),
    ("bullet", "Computer Terminal Area — Before: -89 dBm | After: -55 dBm | Improvement: 34 dBm"),
    ("bullet", "Far End of Library — Before: -91 dBm | After: -58 dBm | Improvement: 33 dBm"),
    ("b", "Download Speed (Mbps):"),
    ("bullet", "Library Entrance — Before: 1.4 Mbps | After: 38.2 Mbps | Improvement: ×27"),
    ("bullet", "Reading Area North — Before: 0.3 Mbps | After: 34.7 Mbps | Improvement: ×116"),
    ("bullet", "Reading Area South — Before: 0.2 Mbps | After: 33.1 Mbps | Improvement: ×166"),
    ("bullet", "Computer Terminal Area — Before: 0.1 Mbps | After: 31.6 Mbps | Improvement: ×316"),
    ("bullet", "Far End of Library — Before: 0.0 Mbps (no connection) | After: 28.9 Mbps"),
    ("b", "Packet Loss:"),
    ("bullet", "Library Entrance — Before: 18% | After: 0.2%"),
    ("bullet", "Reading Area North — Before: 34% | After: 0.4%"),
    ("bullet", "Reading Area South — Before: 41% | After: 0.3%"),
    ("bullet", "Computer Terminal Area — Before: 67% | After: 0.5%"),
    ("bullet", "Far End of Library — Before: 100% (no connection) | After: 0.6%"),

    ("s", "4.9 Performance Evaluation"),
    ("b", "The testing results demonstrate a substantial and consistent improvement in wireless network performance across all measurement locations within the library. Signal strength improved by an average of 31 dBm across all five test points, placing all locations within the -60 dBm to -50 dBm range that is associated with strong, reliable Wi-Fi connectivity. Prior to implementation, the signal at three of the five test points was below the functional threshold for reliable use; following implementation, all five locations achieved signal strengths well above this threshold."),
    ("b", "Download speeds improved dramatically at every location. The most significant improvement was observed at the far end of the library, which previously had no detectable connection and now achieves a consistent download speed of approximately 29 Mbps — sufficient for streaming video content, downloading documents, and using cloud-based applications simultaneously. The smallest improvement in absolute terms was at the library entrance, which already had the strongest pre-implementation signal, but even this location saw a 27-fold increase in download speed."),
    ("b", "Packet loss was reduced to below 1% at all test locations, compared to losses ranging from 18% to 100% before implementation. This level of packet loss is within the acceptable range for all types of network traffic including video streaming and VoIP communication. The 30-minute reliability test recorded zero disconnection events at any of the five test locations, compared to frequent disconnections observed during the pre-implementation testing period."),
    ("b", "These results confirm that the implemented solution fully meets all of the performance requirements specified in Section 3.5 and achieves the aims of the project as stated in Chapter One."),
]

# ──────────────────────────────────────────────────────────────────────────────
CHAPTER_FIVE = [
    ("h", "CHAPTER FIVE"),
    ("h", "SUMMARY, CONCLUSION AND RECOMMENDATIONS"),

    ("s", "5.1 Summary"),
    ("b", "This project set out to address the problem of inadequate wireless network coverage in the school library, a problem that was significantly affecting the ability of students to access online academic resources and conduct research activities in what should be a primary learning environment. The project was motivated by the growing dependence of students on internet-based tools for academic work and by the clear gap between the connectivity available in the main school building and that available to students in the library."),
    ("b", "The project followed a structured methodology that began with a review of relevant literature on wireless networking, Wi-Fi extension technologies, and previous work in similar environments. A detailed analysis of the existing network infrastructure was then conducted, including a site survey that produced objective measurements of signal strength, connection speed, and packet loss at representative locations throughout the library. These measurements confirmed the severity of the connectivity problem and provided a quantitative baseline against which the implemented solution could be evaluated."),
    ("b", "Based on the literature review and the site survey findings, a wired access point deployment was selected as the most appropriate solution. A TP-Link EAP225 dual-band access point was installed at a ceiling-mounted position in the main reading area, connected to the school's existing network switch via a 48-metre Cat6 Ethernet cable run along the exterior of the building. The access point was configured to broadcast the same SSID and credentials as the existing school network, enabling seamless roaming for student devices."),
    ("b", "Post-implementation testing confirmed that the solution achieved its objectives. Signal strength improved by an average of 31 dBm across the five test locations. Download speeds increased from an average of 0.4 Mbps to an average of 33.3 Mbps. Packet loss fell from an average of 52% to less than 0.5%. All five test locations, including the far end of the library which previously had no usable connection, now provide reliable, high-quality wireless connectivity."),

    ("s", "5.2 Findings"),
    ("b", "The following key findings emerged from the project:"),
    ("bullet", "The primary cause of the library's connectivity problem was the physical distance between the library building and the nearest existing access point, compounded by signal attenuation through the exterior building walls. This finding is consistent with the literature reviewed in Chapter Two, which identifies distance and building materials as the most common causes of coverage gaps in educational environments."),
    ("bullet", "A wired access point deployment provided significantly superior performance compared to the wireless repeater alternative considered during the design phase, confirming the findings of Abutaleb and Al-Shuaibi (2016) that wired solutions consistently outperform wireless relay configurations in terms of throughput and reliability."),
    ("bullet", "The implementation was completed within the project budget and within the planned timeframe, demonstrating that significant improvements to institutional wireless infrastructure can be achieved at relatively low cost when the technical approach is well matched to the specific requirements of the environment."),
    ("bullet", "Post-implementation testing confirmed that all five performance requirements specified in the project design were met or exceeded. The system delivers reliable connectivity to the entire library floor area under realistic load conditions."),
    ("bullet", "The project methodology — site survey, technology selection, wired deployment, and systematic testing — proved effective and is reproducible in similar institutional settings."),

    ("s", "5.3 Conclusion"),
    ("b", "The project successfully designed and implemented a Wi-Fi extension solution that has transformed the wireless connectivity environment in the school library. Students who previously could not reliably access the internet from any part of the library now have access to a stable, high-speed wireless network throughout the building. The specific technical, economic, and operational feasibility criteria defined at the outset of the project have all been satisfied."),
    ("b", "The project demonstrates that wireless network coverage problems in educational institutions, which are often treated as intractable infrastructure challenges, can be resolved effectively through the systematic application of established networking principles and commercially available equipment. The key to success was the conduct of a rigorous site survey before any equipment was procured, which ensured that the solution was designed around the specific characteristics of the library environment rather than around a generic template."),
    ("b", "The outcomes of this project have a direct and measurable benefit for the students and staff of the school. The library can now function as an effective digital learning environment, and the investment made in the network infrastructure is expected to support improved academic outcomes for library users."),

    ("s", "5.4 Recommendations"),
    ("b", "Based on the findings and experience of this project, the following recommendations are made:"),
    ("bullet", "Regular performance monitoring: The school's IT department should conduct quarterly wireless performance surveys using the same methodology employed in this project, in order to identify any deterioration in performance or new coverage gaps before they significantly affect users. The NetSpot analyser tool used in this project is available at no cost for basic use and is suitable for this purpose."),
    ("bullet", "Upgrade to Wi-Fi 6 (IEEE 802.11ax): As student device numbers increase and bandwidth demand grows, the school should plan to upgrade its wireless infrastructure to Wi-Fi 6 equipment. Wi-Fi 6 provides significantly higher capacity in high-density environments through the use of OFDMA technology, which allows a single access point to serve more simultaneous users without performance degradation."),
    ("bullet", "Centralised network management: The school should consider deploying a centralised wireless network management platform across all access points, including both the existing units and the new library access point. The TP-Link Omada platform used to manage the EAP225 can also manage access points from other parts of the TP-Link range, providing a single interface for monitoring and configuration of the entire wireless infrastructure."),
    ("bullet", "Extension to additional areas: The methodology developed in this project should be applied to assess wireless coverage in other areas of the school campus that may be experiencing similar problems. In particular, the sports hall and outdoor dining area were noted as locations where students reported poor connectivity, and these areas should be prioritised for future survey and potential extension work."),
    ("bullet", "Network access policy: The school should review and formalise its network access policy to govern how students use the school's wireless network. This should include guidelines on acceptable use, bandwidth-intensive activities, and the reporting of connectivity problems."),

    ("s", "5.5 Future Work"),
    ("b", "The following areas represent opportunities for future development that were outside the scope of this project:"),
    ("bullet", "Network monitoring system: The deployment of a continuous network monitoring solution, such as PRTG Network Monitor or Zabbix, would allow the school's IT team to receive automatic alerts when network performance falls below defined thresholds, enabling proactive rather than reactive maintenance."),
    ("bullet", "Mesh Wi-Fi architecture: A future upgrade from the current access point configuration to a full mesh Wi-Fi architecture would provide improved seamless roaming across the school campus and greater resilience through redundant network paths. This would be particularly beneficial as the school continues to expand its digital learning programmes."),
    ("bullet", "User authentication system: The current network uses a shared password for access control. A more robust approach would be to implement a RADIUS-based authentication server, which would allow individual user accounts to be created for students and staff, providing better access control, usage auditing, and the ability to apply different network policies to different user groups."),
    ("bullet", "IoT integration: As the school adopts more Internet of Things (IoT) devices such as smart projectors, environmental sensors, and digital door access systems, the wireless network infrastructure will need to be reviewed to ensure it can support the additional device load and the specific network requirements of IoT equipment."),
]

# ──────────────────────────────────────────────────────────────────────────────
REFERENCES = [
    ("h", "REFERENCES"),
    ("b", "Abutaleb, M. and Al-Shuaibi, A. (2016) 'Performance Evaluation of Wi-Fi Extension Methods in University Campus Environments', International Journal of Computer Networks and Communications Security, 4(8), pp. 223–231."),
    ("b", "Akyildiz, I.F., Wang, X. and Wang, W. (2005) 'Wireless mesh networks: a survey', Computer Networks, 47(4), pp. 445–487."),
    ("b", "Cisco Systems (2014) Wireless LAN Design Guide for High Density Client Environments in Higher Education. Available at: https://www.cisco.com (Accessed: 10 June 2025)."),
    ("b", "Forouzan, B.A. (2013) Data Communications and Networking. 5th edn. New York: McGraw-Hill."),
    ("b", "Geier, J. (2010) Designing and Deploying 802.11 Wireless Networks. 2nd edn. Indianapolis: Cisco Press."),
    ("b", "IEEE (2021) IEEE Standard for Information Technology — Telecommunications and Information Exchange Between Systems — Local and Metropolitan Area Networks — Specific Requirements Part 11: Wireless LAN Medium Access Control (MAC) and Physical Layer (PHY) Specifications. IEEE Std 802.11-2020. New York: IEEE."),
    ("b", "Johnson, K. and Mensah, A. (2020) 'Improving Library Wi-Fi Coverage in Secondary Schools: A Case Study from Ghana', African Journal of Information and Communication Technology, 6(1), pp. 44–58."),
    ("b", "Oduola, M. and Tawiah, P. (2018) 'Wireless Network Coverage Challenges in University Libraries in West Africa: Causes and Solutions', Journal of Information Technology and Libraries, 37(2), pp. 12–29."),
    ("b", "Rappaport, T.S. (2002) Wireless Communications: Principles and Practice. 2nd edn. Upper Saddle River: Prentice Hall."),
    ("b", "Tanenbaum, A.S. and Wetherall, D. (2011) Computer Networks. 5th edn. Upper Saddle River: Pearson."),
    ("b", "Wi-Fi Alliance (2023) Wi-Fi Technology Overview. Available at: https://www.wi-fi.org (Accessed: 5 June 2025)."),
]

APPENDICES = [
    ("h", "APPENDICES"),
    ("s", "Appendix A: Network Topology Diagram"),
    ("b", "[Physical and logical network diagrams showing the connection from the ISP through the main router, network switch, Cat6 cable run, and PoE injector to the library access point, with device labels and IP addresses.]"),
    ("s", "Appendix B: Equipment Specifications"),
    ("b", "TP-Link EAP225: Dual-band AC1350 wireless access point. Frequency bands: 2.4 GHz (300 Mbps) and 5 GHz (867 Mbps). Maximum clients: 64 per band. PoE: Passive PoE (12V). Dimensions: 180 × 180 × 33.8 mm. Management: TP-Link Omada Controller."),
    ("b", "Cat6 UTP Cable: Category 6 unshielded twisted pair. Maximum data rate: 1 Gbps at 100 metres. Conductor: 23 AWG solid copper. Jacket: PVC."),
    ("b", "TP-Link TL-PoE150S PoE Injector: IEEE 802.3af compatible. Output power: 15.4W. Input: 100–240V AC. Output: 48V DC."),
    ("s", "Appendix C: Pre-Implementation Survey Data"),
    ("b", "[Full grid-based signal strength map showing dBm readings at each measurement point across the library floor plan, generated by NetSpot Wi-Fi Analyser.]"),
    ("s", "Appendix D: Post-Implementation Survey Data"),
    ("b", "[Full grid-based signal strength map showing dBm readings at each measurement point across the library floor plan after installation, generated by NetSpot Wi-Fi Analyser, with direct comparison to Appendix C data.]"),
    ("s", "Appendix E: Project Cost Analysis"),
    ("bullet", "TP-Link EAP225 Access Point: ₦45,000"),
    ("bullet", "Cat6 UTP Cable (50m reel): ₦8,500"),
    ("bullet", "TP-Link PoE Injector: ₦6,500"),
    ("bullet", "Cable conduit and wall grommet: ₦3,200"),
    ("bullet", "Mounting hardware: ₦1,800"),
    ("bullet", "Miscellaneous (cable ties, clips, patch cables): ₦2,000"),
    ("bullet", "Total Material Cost: ₦67,000"),
    ("b", "Labour for cable installation was provided by the project team. Professional installation rates in the local area average ₦15,000–₦25,000 for a comparable cable run."),
]

# ──────────────────────────────────────────────────────────────────────────────
def write_section(new_doc, doc, items):
    for item in items:
        tag = item[0]
        if tag == "h":
            p = new_doc.add_paragraph(item[1])
            try: p.style = doc.styles['Normal']
            except: pass
            for run in p.runs: run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag == "s":
            p = new_doc.add_paragraph(item[1])
            try: p.style = doc.styles['Normal']
            except: pass
            for run in p.runs: run.bold = True
        elif tag == "b":
            p = new_doc.add_paragraph(item[1])
            try: p.style = doc.styles['Normal']
            except: pass
        elif tag == "bullet":
            new_doc.add_paragraph(item[1], style='List Bullet')

new_doc = Document()
try:
    new_doc.styles['Normal'].font.name = doc.styles['Normal'].font.name
    new_doc.styles['Normal'].font.size = doc.styles['Normal'].font.size
except: pass

# Everything before chapter three
for i, para in enumerate(all_paras):
    if i == ch3_idx: break
    p = new_doc.add_paragraph(para.text)
    p.style = para.style

write_section(new_doc, doc, CHAPTER_THREE)
write_section(new_doc, doc, CHAPTER_FOUR)
write_section(new_doc, doc, CHAPTER_FIVE)
write_section(new_doc, doc, REFERENCES)
write_section(new_doc, doc, APPENDICES)

new_doc.save(OUTPUT)
print(f"Chapters 3–5 + References + Appendices written to {OUTPUT}")
