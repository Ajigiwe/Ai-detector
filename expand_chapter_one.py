from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# Read the original document
doc = Document(r'docs/final year project for networking student (Autosaved).docx')
all_paras = list(doc.paragraphs)

# Find chapter one start and chapter two start indices in the actual doc
ch1_start = None
ch2_start = None
for i, p in enumerate(all_paras):
    if p.text.strip() == "CHAPTER ONE":
        ch1_start = i
    if p.text.strip() == "CHAPTER TWO":
        ch2_start = i
        break

print(f"Chapter One starts at paragraph index: {ch1_start}")
print(f"Chapter Two starts at paragraph index: {ch2_start}")

# The expanded Chapter One content
CHAPTER_ONE_EXPANDED = [
    ("heading", "CHAPTER ONE"),
    ("heading", "INTRODUCTION"),

    ("section", "1.1 Background of the Study"),
    ("body", "The internet has become an indispensable resource in modern educational institutions worldwide. Students today depend heavily on online platforms for academic research, assignment submission, access to digital libraries, e-learning portals, and communication with peers and lecturers. This growing reliance on internet-based learning tools has made reliable wireless network connectivity a critical component of any academic environment."),
    ("body", "Wireless networking technology, particularly Wi-Fi, has transformed the way educational institutions deliver internet access. Unlike traditional wired networks that require fixed cable installations, Wi-Fi allows students and staff to connect their devices freely within the coverage area. This flexibility has made wireless networking the preferred standard for internet access in schools, universities, and libraries across the world."),
    ("body", "A school library occupies a unique and important position within an academic institution. It serves as the primary space where students engage in independent research, group study, reading, and the use of digital resources. For this reason, consistent and reliable internet access within the library is not merely a convenience but a necessity for effective learning. Students who visit the library expect to be able to connect to the internet without interruption in order to access journals, databases, e-books, and academic resources."),
    ("body", "However, many school libraries suffer from inadequate wireless network coverage. This is commonly caused by the physical distance of the library from the main network infrastructure, the interference created by thick walls and structural materials, the limitations of the primary router's signal range, and the increasing number of devices competing for the same network bandwidth. These factors combine to create dead zones — areas within the library where the Wi-Fi signal is either too weak to be useful or entirely absent."),
    ("body", "The consequence of poor wireless coverage in a library is significant. Students are unable to complete research tasks, access required academic materials, submit assignments online, or participate in digital learning activities. This directly undermines the educational purpose of the library and places students at a disadvantage compared to those with reliable internet access. In an institution where academic performance is closely linked to access to digital resources, connectivity gaps of this kind cannot be overlooked."),
    ("body", "This project seeks to address these challenges by conducting a systematic analysis of the existing wireless network infrastructure in the school library, identifying specific areas of weak or absent signal coverage, and designing and implementing a Wi-Fi extension solution that significantly improves wireless coverage and internet accessibility throughout the library. The outcome of this project is intended to benefit both students and staff by providing a stable, accessible, and well-performing wireless network environment that supports the full range of academic activities."),

    ("section", "1.2 Problem Statement"),
    ("body", "Despite the growing importance of internet access in academic learning, students using the school library face persistent and significant challenges in connecting to the wireless network. The existing network infrastructure, which was designed to serve the main administrative and classroom areas of the school, does not adequately extend to cover the library building. This has resulted in a range of connectivity problems that directly affect the ability of students to use the library as an effective learning environment."),
    ("body", "The specific problems experienced by library users include the following:"),
    ("bullet", "Weak Wi-Fi signal strength: The signal from the main router deteriorates significantly by the time it reaches the library, resulting in a low signal level that is often insufficient for reliable internet use."),
    ("bullet", "Slow internet connection speeds: Even in areas where a signal is detectable, the connection speed is too slow to support normal academic activities such as streaming educational content, downloading documents, or using cloud-based applications."),
    ("bullet", "Frequent disconnections: Students experience repeated drops in connection, requiring them to reconnect manually. This disrupts research sessions and reduces productivity."),
    ("bullet", "Limited network coverage area: Large portions of the library have no detectable Wi-Fi signal at all, meaning students in those areas cannot connect regardless of their device capabilities."),
    ("bullet", "Difficulty accessing online academic resources: The unreliable connection prevents students from consistently accessing online journals, digital libraries, institutional portals, and other resources that are essential for academic work."),
    ("body", "These problems collectively reduce the productivity of students who use the library, limit the educational value of the library as a research space, and create inequitable access to digital learning resources. Without a targeted technical intervention, the library will continue to underperform as an academic facility. This project is therefore necessary to identify the root causes of the connectivity problems and implement an appropriate and cost-effective solution."),

    ("section", "1.3 Aim of the Project"),
    ("body", "The primary aim of this project is to design and implement a Wi-Fi extension system that significantly improves wireless internet coverage, reliability, and accessibility throughout the school library. The project will identify the specific technical shortcomings of the existing wireless infrastructure as it relates to the library environment, propose a suitable extension solution based on established networking principles, and validate the effectiveness of the solution through systematic testing and performance evaluation."),
    ("body", "The aim is not simply to add additional networking hardware but to deliver a measurable and sustained improvement in wireless connectivity that enables students to carry out their academic activities without interruption. The solution must be practical, appropriately scaled for the library environment, and achievable within the constraints of the project's available resources."),

    ("section", "1.4 Objectives"),
    ("body", "In order to achieve the stated aim of the project, the following specific objectives have been defined:"),
    ("bullet", "To analyse the existing wireless network infrastructure of the school and assess its current performance within the library environment, including signal strength, coverage area, and connection speeds."),
    ("bullet", "To identify and map specific areas within the library that experience weak Wi-Fi signal or no coverage, using appropriate tools and measurement methods."),
    ("bullet", "To evaluate suitable Wi-Fi extension technologies and select the most appropriate solution based on the requirements of the library environment, available budget, and technical feasibility."),
    ("bullet", "To design a wireless network extension plan that specifies the placement of additional access points or extenders, the required equipment, cabling routes, IP addressing scheme, and security configuration."),
    ("bullet", "To implement the designed Wi-Fi extension system by procuring, installing, and configuring the necessary hardware and software components."),
    ("bullet", "To test the implemented network by measuring signal strength, coverage area, connection speed, and reliability across the library before and after implementation."),
    ("bullet", "To evaluate the overall effectiveness of the solution by comparing pre-implementation and post-implementation performance data and documenting the improvements achieved."),

    ("section", "1.5 Research Questions"),
    ("body", "This project is guided by the following research questions, which define the scope of investigation and provide a framework for analysis:"),
    ("bullet", "What are the current wireless network conditions within the school library, and what specific factors are responsible for the inadequate Wi-Fi coverage experienced by students?"),
    ("bullet", "Which Wi-Fi extension technologies and methods are most suitable for addressing the connectivity challenges identified in the library environment, considering factors such as cost, ease of implementation, and performance?"),
    ("bullet", "What equipment, configuration, and network design are required to effectively implement the chosen Wi-Fi extension solution in the school library?"),
    ("bullet", "To what extent does the implemented Wi-Fi extension system improve wireless coverage, signal strength, connection speed, and overall internet accessibility within the library compared to the previous infrastructure?"),

    ("section", "1.6 Significance of Study"),
    ("body", "This project holds considerable significance for the school and its students, as well as for the broader field of educational network infrastructure. The following points outline the key areas in which the project is expected to make a meaningful contribution:"),
    ("bullet", "Improved student learning outcomes: By providing reliable wireless internet access throughout the library, students will be able to access digital academic resources, complete research assignments, and engage with online learning platforms without the disruption caused by poor connectivity. This is expected to have a direct positive impact on the quality and productivity of academic work."),
    ("bullet", "Enhanced library utility: The library will be able to fully realise its function as a centre for research and learning. Students will be more likely to use the library as an effective study space when they can trust that internet access will be available and consistent."),
    ("bullet", "Equitable access to digital resources: Students who rely on the school's network infrastructure as their primary means of internet access will benefit most from this improvement. The project contributes to reducing the digital divide within the school by ensuring that all students, regardless of where they sit within the library, have access to the same quality of connectivity."),
    ("bullet", "Practical application of networking knowledge: This project demonstrates the practical application of wireless networking principles, network analysis, equipment configuration, and performance testing in a real-world educational environment. It serves as a model for how networking projects can be designed and executed in similar institutional settings."),
    ("bullet", "Template for future network expansion: The methodology, design decisions, and implementation approach documented in this project can serve as a reference for future network improvement projects within the same school or in similar institutions facing comparable connectivity challenges."),

    ("section", "1.7 Scope of Study"),
    ("body", "The scope of this project is defined and bounded by the following parameters:"),
    ("body", "The project is limited to the wireless network infrastructure of the school library building. It does not extend to other areas of the school campus such as classrooms, administrative offices, laboratories, or outdoor spaces. The investigation and proposed solution are specifically targeted at improving Wi-Fi coverage and performance within the library environment."),
    ("body", "The study covers the analysis of the existing wireless network as it pertains to library coverage, the design and implementation of an appropriate extension solution, and the testing and evaluation of the implemented system. It does not include a redesign of the school's core network infrastructure, changes to the internet service provider's equipment, or modifications to the school's main router configuration beyond what is necessary to support the extension."),
    ("body", "The project is conducted within the timeframe of the academic year and with the resources available to the project group. Recommendations that fall outside the current project scope — such as institution-wide network upgrades — are noted in the recommendations chapter but are not implemented as part of this work."),

    ("section", "1.8 Limitations"),
    ("body", "The following limitations were identified during the planning and execution of this project:"),
    ("bullet", "Budget constraints: The project was carried out with a limited budget, which restricted the range of equipment that could be considered. Higher-performance enterprise-grade access points and professional installation services were not feasible within the available resources. The equipment selected represents the best available option within the budget constraints."),
    ("bullet", "Equipment availability: Some networking equipment considered during the design phase was not readily available locally and could not be obtained within the project timeline. Substitutions were made where necessary, and the impact of these substitutions on performance is discussed in the relevant chapters."),
    ("bullet", "Time limitations: The project was completed within the constraints of an academic calendar, which limited the amount of time available for extended testing and long-term performance monitoring. Post-implementation testing was conducted over a defined period, and the results reflect the performance of the system within that window."),
    ("bullet", "Access restrictions: Physical access to certain parts of the library building, including ceiling spaces and wall cavities required for optimal cable routing, was limited. This influenced some of the installation decisions and equipment placement choices made during implementation."),
    ("bullet", "User participation: The evaluation of user satisfaction was based on a sample of available library users during the testing period and may not fully represent the experience of all student users across different times of day and usage patterns."),

    ("section", "1.9 Definition of Terms"),
    ("body", "The following terms are used throughout this project and are defined here for clarity:"),
    ("term", "Wi-Fi", "A wireless networking technology that uses radio frequency signals to provide internet access and local network connectivity to devices within a defined coverage area. Wi-Fi operates according to the IEEE 802.11 family of standards and is the primary means of wireless internet access in most modern environments."),
    ("term", "Access Point (AP)", "A hardware device that creates a wireless local area network (WLAN) within a building or area. An access point connects to a wired network via an Ethernet cable and broadcasts a wireless signal that allows Wi-Fi-enabled devices to connect to the network."),
    ("term", "Router", "A networking device that receives, processes, and forwards data packets between computer networks. In a typical home or institutional network, the router connects the local network to the internet and manages traffic between devices on the local network."),
    ("term", "Wi-Fi Extender / Range Extender", "A device that receives an existing Wi-Fi signal, amplifies it, and retransmits it to extend the coverage area of the wireless network. A range extender can help eliminate dead zones in areas where the original router signal is too weak to be reliable."),
    ("term", "SSID (Service Set Identifier)", "The name assigned to a wireless network, which is broadcast by the access point or router and displayed to users when they search for available Wi-Fi networks. Devices use the SSID to identify and connect to the correct network."),
    ("term", "Bandwidth", "The maximum rate at which data can be transmitted over a network connection, typically measured in megabits per second (Mbps) or gigabits per second (Gbps). Higher bandwidth supports faster data transfer and can accommodate more simultaneous users."),
    ("term", "Signal Strength (RSSI)", "A measure of the power level of a wireless signal received by a device, typically expressed in decibels relative to a milliwatt (dBm). Signal strength indicates how well a device can communicate with the access point; stronger signals generally result in faster and more reliable connections."),
    ("term", "Dead Zone", "An area within a building or environment where the wireless network signal is too weak to support a usable connection. Dead zones are commonly caused by physical obstructions, distance from the access point, or interference from other electronic devices."),
    ("term", "IEEE 802.11", "A set of technical standards developed by the Institute of Electrical and Electronics Engineers (IEEE) that define the protocols for implementing wireless local area network (WLAN) communication. Common variants include 802.11b, 802.11g, 802.11n, 802.11ac, and 802.11ax (Wi-Fi 6)."),
    ("term", "Network Topology", "The physical or logical arrangement of devices and connections in a computer network. Common topologies include star, bus, ring, and mesh configurations, each with different implications for performance, reliability, and scalability."),
]

# Build the new document by replacing chapter one content
new_doc = Document()

# Copy styles from original if possible
try:
    new_doc.styles['Normal'].font.name = doc.styles['Normal'].font.name
    new_doc.styles['Normal'].font.size = doc.styles['Normal'].font.size
except:
    pass

# Write paragraphs before chapter one (title page, abstract, TOC)
for i, para in enumerate(all_paras):
    if i == ch1_start:
        break
    p = new_doc.add_paragraph(para.text)
    p.style = para.style

# Write expanded chapter one
for item in CHAPTER_ONE_EXPANDED:
    if item[0] == "heading":
        p = new_doc.add_paragraph(item[1])
        try:
            p.style = doc.styles['Normal']
        except:
            pass
        for run in p.runs:
            run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif item[0] == "section":
        p = new_doc.add_paragraph(item[1])
        try:
            p.style = doc.styles['Normal']
        except:
            pass
        for run in p.runs:
            run.bold = True

    elif item[0] == "body":
        p = new_doc.add_paragraph(item[1])
        try:
            p.style = doc.styles['Normal']
        except:
            pass

    elif item[0] == "bullet":
        p = new_doc.add_paragraph(item[1], style='List Bullet')

    elif item[0] == "term":
        # Bold term name, then colon, then definition
        p = new_doc.add_paragraph()
        run_term = p.add_run(item[1] + ": ")
        run_term.bold = True
        run_def = p.add_run(item[2])

# Write all paragraphs from chapter two onwards
for i in range(ch2_start, len(all_paras)):
    p = new_doc.add_paragraph(all_paras[i].text)
    p.style = all_paras[i].style

new_doc.save(r'docs/final year project for networking student (Autosaved).docx')
print("Chapter One expanded successfully.")
print(f"New document saved.")
